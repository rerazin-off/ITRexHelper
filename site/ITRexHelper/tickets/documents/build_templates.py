"""Создание DOCX-шаблонов для модуля импортирования документов."""

from pathlib import Path

from django.conf import settings
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BASE_DIR = Path(settings.BASE_DIR)
TEMPLATE_DIR = BASE_DIR / 'document_templates'


def _add_heading(doc, text, size=16):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return paragraph


def create_contract_template():
    doc = Document()
    _add_heading(doc, 'ДОГОВОР ОКАЗАНИЯ УСЛУГ № {{ contract_number }}')

    paragraph = doc.add_paragraph('г. Нижний Новгород\t\t{{ contract_date }}')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        '{{ contractor_name }}, именуемое в дальнейшем «Исполнитель», '
        'и {{ client_company }} в лице {{ client_name }}, '
        'именуемый(ая) в дальнейшем «Заказчик», заключили настоящий договор о нижеследующем:'
    )

    doc.add_paragraph('1. Предмет договора').runs[0].bold = True
    doc.add_paragraph(
        '1.1. Исполнитель обязуется оказать услуги технической поддержки по заявке '
        '№ {{ ticket_id }} от {{ ticket_date }}: {{ ticket_description }}.'
    )

    doc.add_paragraph('2. Стороны договора').runs[0].bold = True
    doc.add_paragraph('2.1. Исполнитель: {{ contractor_name }}, {{ contractor_details }}.')
    doc.add_paragraph('2.2. Заказчик: {{ client_company }}, контакт: {{ client_contact }}.')

    doc.add_paragraph('3. Сроки и порядок исполнения').runs[0].bold = True
    doc.add_paragraph(
        '3.1. Услуги оказываются в рамках заявки № {{ ticket_id }}. '
        'Статус заявки: {{ ticket_status }}.'
    )

    doc.add_paragraph('4. Подписи сторон').runs[0].bold = True
    doc.add_paragraph('')
    doc.add_paragraph('Исполнитель: {{ contractor_name }}')
    doc.add_paragraph('________________ / ITREX /')
    doc.add_paragraph('')
    doc.add_paragraph('Заказчик: {{ client_company }}')
    doc.add_paragraph('________________ / {{ client_name }} /')

    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(TEMPLATE_DIR / 'contract.docx')


def create_analytics_template():
    doc = Document()
    _add_heading(doc, 'IT REX HELPER', 14)
    _add_heading(doc, 'Отчёт по аналитике заявок')

    paragraph = doc.add_paragraph('Дата формирования: {{ report_date }}')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Сводные показатели').runs[0].bold = True
    table = doc.add_table(rows=6, cols=2)
    rows = [
        ('Показатель', 'Значение'),
        ('Всего заявок', '{{ total_count }}'),
        ('В работе', '{{ in_progress_count }}'),
        ('Просрочено', '{{ overdue_count }}'),
        ('Заявок в этом месяце', '{{ tickets_month }}'),
        ('Обработано за месяц', '{{ processed_month }}'),
    ]
    for index, (label, value) in enumerate(rows):
        table.rows[index].cells[0].text = label
        table.rows[index].cells[1].text = value

    doc.add_paragraph('Динамика за последние 6 месяцев').runs[0].bold = True
    doc.add_paragraph('{% for month in chart_months %}{{ month }}{% if not loop.last %}, {% endif %}{% endfor %}')

    doc.add_paragraph('Журнал аудита (последние изменения)').runs[0].bold = True
    doc.add_paragraph(
        '{% for row in audit_rows %}'
        'Заявка {{ row.ticket }}: {{ row.old_status }} → {{ row.new_status }} ({{ row.changed_at }})\n'
        '{% endfor %}'
    )

    doc.save(TEMPLATE_DIR / 'analytics_report.docx')


if __name__ == '__main__':
    create_contract_template()
    create_analytics_template()
    print('Templates created in', TEMPLATE_DIR)
